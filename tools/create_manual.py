from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\Daiene\Documents\Codex\2026-08-16\eu\docs\Manual-RotaOS-Camilla.docx")
PURPLE = "6E55D7"
INK = "202337"
MUTED = "6F7483"
SOFT = "F0EDFF"
GREEN = "27895F"
RED = "A13F48"


def set_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=INK):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run, 9.5, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(text)


def add_callout(doc, title, text, fill=SOFT, title_color=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, 10, True, title_color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), 9.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        shade(cell, "E8EEF5")
        cell_text(cell, header, True, "1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [("Heading 1", 16, "2E74B5", 14, 7), ("Heading 2", 13, "2E74B5", 11, 5), ("Heading 3", 12, "1F4D78", 8, 4)]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("RotaOS | Manual de apresentacao"), 8.5, True, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("RotaOS - Planejamento de rotas"), 8.5, False, MUTED)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(50)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("ROTAOS"), 13, True, PURPLE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_font(p.add_run("Manual de funcionamento"), 28, True, INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(34)
set_font(p.add_run("Planejamento de rotas para chamados do Procesa"), 14, False, MUTED)
add_callout(doc, "Objetivo da apresentacao", "Explicar como o RotaOS recebe as OS ativas, ajuda a Camilla a priorizar e distribui sugestoes para revisao. Nenhuma acao e enviada automaticamente ao Procesa.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Versao de referencia: agosto de 2026"), 9.5, False, MUTED)

doc.add_page_break()
add_heading(doc, "1. O que o RotaOS faz")
add_body(doc, "O RotaOS e uma camada de planejamento. Ele organiza os chamados ativos da aba Solicitadas do Procesa, registra a base online e apresenta uma sugestao de prioridades e de distribuicao para a Camilla conferir.")
add_table(doc, ["Etapa", "Resultado"], [
    ("1. Coleta", "A extensao le as OS da aba Solicitadas no Procesa."),
    ("2. Filtro", "Somente S025, S200 X e S201 X entram no RotaOS."),
    ("3. Banco", "A base online guarda e atualiza os chamados importados."),
    ("4. Analise", "O RotaOS calcula prioridades e monta sugestoes revisaveis."),
    ("5. Revisao", "Camilla escolhe, ajusta e usa a sugestao como apoio."),
], [1.25, 5.25])
add_callout(doc, "Regra de seguranca", "O RotaOS nao envia distribuicoes ao Procesa. A distribuicao mostrada no painel e apenas uma proposta para revisao.", "EAF7F1", GREEN)
add_heading(doc, "2. Como atualizar a base", 1)
for text in [
    "Entrar no RotaOS com o usuario cadastrado.",
    "Abrir o Procesa na aba Solicitadas. O filtro visual e opcional: a extensao filtra S025, S200 e S201 mesmo quando o campo Servico Solicitado nao esta filtrado.",
    "Usar Todas as paginas de Solicitadas e aguardar a contagem final da coleta.",
    "Voltar ao RotaOS, que salva, compara e recarrega a base online.",
]: add_bullet(doc, text)
add_callout(doc, "Como validar a coleta", "Compare o total final informado pela extensao com o total de OS compatveis na aba Solicitadas. A base do RotaOS pode ter menos itens se o Procesa trouxer outros servicos alem de S025, S200 e S201.", "FFF7DE", "8A6315")

add_heading(doc, "3. Regras que definem a prioridade", 1)
add_body(doc, "A sugestao e uma lista de apoio. Ela nao marca nem desmarca a selecao manual da Camilla.")
add_table(doc, ["Regra", "Como funciona hoje"], [
    ("Reclamacoes recentes", "A quantidade de reclamacoes soma pontos somente quando a reclamacao mais recente ocorreu nos ultimos 30 dias. A pontuacao considera no maximo cinco reclamacoes para evitar distorcao."),
    ("Recencia da solicitacao", "Chamados mais novos recebem mais pontos. A pontuacao diminui gradualmente com o tempo."),
    ("Criterios adicionais", "Camilla pode adicionar pontos por bairro, regiao ou servico no menu Criterios da sugestao."),
    ("Compatibilidade", "Uma equipe so recebe OS dos servicos que ela atende e dentro da capacidade diaria configurada."),
    ("Mesmo endereco", "OS do mesmo endereco sao avaliadas juntas para favorecer uma rota mais coerente."),
], [1.65, 4.85])
add_callout(doc, "Exemplo importante", "Uma OS antiga pode continuar no topo se recebeu reclamacoes recentemente. A idade da OS nao apaga a urgencia de uma nova reclamacao.")

doc.add_page_break()
add_heading(doc, "4. Equipes: RotaOS x Procesa")
add_body(doc, "As equipes configuradas no RotaOS sao as equipes disponiveis para a sugestao. O nome que eventualmente aparece na coluna Equipe do Procesa e uma informacao historica da OS e nao vira uma atribuicao automatica no RotaOS.")
add_table(doc, ["Informacao", "Uso no RotaOS"], [
    ("Equipe cadastrada no RotaOS", "Define servicos atendidos, capacidade diaria e participa da sugestao."),
    ("Equipe exibida no Procesa", "Sera sinalizada para conferencia. Pode indicar tentativa anterior de distribuicao ainda sem finalizacao."),
    ("OS sem equipe compativel", "Permanece sem atribuicao na sugestao. Nada e forcado para uma equipe inadequada."),
], [1.85, 4.65])
add_callout(doc, "Decisao a validar com a Camilla", "Quando uma OS ainda esta em Solicitadas, mas ja mostra equipe no Procesa: devemos manter a mesma equipe, ou escolher a equipe mais proxima da rota planejada? A resposta vira uma regra futura do RotaOS.", "FFF0F1", RED)
add_heading(doc, "5. Como ler a Base de chamados", 1)
for text in [
    "Atualizacao: indica se a OS e nova, alterada, conferida ou possui reclamacao.",
    "Reclamacoes: mostra a quantidade coletada na impressao do Procesa e a data da ultima ocorrencia, quando disponivel.",
    "Detalhes: mostra o relato mais recente da reclamacao; ele ajuda a entender a urgencia.",
    "Sugestao RotaOS: aparece somente depois de gerar uma sugestao. Ver motivo explica a pontuacao.",
    "Equipe: mostra a equipe sugerida no RotaOS; A definir significa que nenhuma equipe cadastrada era compativel ou tinha capacidade.",
]: add_bullet(doc, text)
add_callout(doc, "Tela mais compacta", "O modo padrao prioriza a Base de chamados. Use Mostrar visao completa quando quiser aumentar o mapa e o painel de equipes.")

add_heading(doc, "6. Mapa de Juiz de Fora: proxima etapa", 1)
add_body(doc, "Sim, o mapa pode ser um mapa real de Juiz de Fora, MG. Para ele orientar rotas de verdade, cada endereco precisa ganhar coordenadas geograficas e essas coordenadas precisam ser guardadas no banco.")
add_table(doc, ["Fase", "Entrega"], [
    ("Base geografica", "Mapa real de Juiz de Fora com zoom, ruas e marcadores das OS."),
    ("Geocodificacao", "Conversao de endereco, bairro e cidade em latitude e longitude, com conferencia de enderecos nao encontrados."),
    ("Rotas", "Agrupamento por equipe e estimativa de deslocamento entre OS."),
    ("Aprimoramento", "Regra para priorizar a equipe ja indicada no Procesa ou a equipe mais proxima, conforme decisao da Camilla."),
], [1.6, 4.9])
add_callout(doc, "Limite atual", "O mapa atual e ilustrativo. Ele mostra a ideia de distribuicao, mas ainda nao calcula distancia real nem navega pelas ruas.", "FFF7DE", "8A6315")

doc.add_page_break()
add_heading(doc, "7. Roteiro curto para a apresentacao", 1)
steps = [
    "Abrir a Base de chamados e mostrar que o foco e a leitura das OS ativas.",
    "Explicar a coleta no Procesa e o filtro automatico dos tres servicos atendidos.",
    "Pesquisar uma OS com reclamacao e mostrar quantidade, ultimo relato e motivo da prioridade.",
    "Gerar uma sugestao e explicar que ela nao altera a selecao manual nem envia nada ao Procesa.",
    "Mostrar as equipes, seus servicos e suas capacidades diarias.",
    "Apresentar a pergunta sobre OS que ja possuem equipe exibida no Procesa.",
    "Fechar com a proxima evolucao: mapa real de Juiz de Fora e rotas por proximidade.",
]
for number, text in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)
add_heading(doc, "8. Checklist antes de usar", 1)
for text in [
    "Estou logada no RotaOS?", "A extensao instalada e a versao atual?", "Estou na aba Solicitadas do Procesa?", "A coleta terminou todas as paginas?", "O total parece compativel com os servicos S025, S200 e S201?", "As equipes e capacidades estao atualizadas?", "As OS com equipe exibida no Procesa foram conferidas?"]:
    add_bullet(doc, "[ ] " + text)
add_callout(doc, "Mensagem final", "O RotaOS organiza a decisao da Camilla. A experiencia humana continua no centro: o sistema propoe, a Camilla confere e decide.", "EAF7F1", GREEN)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
